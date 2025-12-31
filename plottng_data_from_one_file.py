import os
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from correlation import corr_stats
from utils import merge_tables, open_file, calculate_soc_changes
from settings import RESULTS
from apsimNGpy.core_utils.database_utils import read_db_table
from apsimNGpy.core.plotmanager import PlotManager
from dataclasses import dataclass
import seaborn as sns

sns.set_style("darkgrid", rc={"font.family": "DejaVu Sans"})
plot_dir = RESULTS / ('Plots')
plot_dir.mkdir(parents=True, exist_ok=True)


# _____________ constants _____________________
@dataclass(frozen=True, slots=True, repr=False, order=False)
class PlotConstants:
    xLABEL_SIZE = 18
    yLABEL_SIZE = 18
    TITLE_SIZE = 20
    xTICK_SIZE = 14
    yTICK_SIZE = 14
    LEGEND_SIZE = 12

    LINE_WIDTH = 2.0
    MARKER_SIZE = 6
    ALPHA = 0.9
    HEIGHT = 8
    ASPECT_RATIO = 1.2
    FIGSIZE = (8, 5)  # inches (width, height)
    DPI = 600
    PALETTE = 'tab10'
    GRID = True
    GRID_STYLE = "--"
    GRID_ALPHA = 0.3
    FILE_EXTENSION = ".png"
    DATE_FMT = "%Y-%m-%d"
    CURRENCY_FMT = "${:,.0f}"
    PCT_FMT = "{:.1%}"

    # Matplotlib color cycle (optional)
    COLOR_CYCLE = [
        "#1f77b4", "#ff7f0e", "#2ca02c", "#d62728",
        "#9467bd", "#8c564b", "#e377c2", "#7f7f7f",
        "#bcbd22", "#17becf"
    ]


LABELS = {
    'Surf_Org_Carbon_Mg': 'Surface organic carbon (Mg ha⁻¹)',
    'cnr': 'C:N ratio ',
    'SurfaceOrganicMatter_Nitrogen': 'Surface organic matter N (kg ha⁻¹)',
    'top_mineralized_N': 'Mineralized nitrogen (0-15cm, kg ha⁻¹)',
    'Residue': 'Incorporated residue fraction',
    'Nitrogen': 'Nitrogen fertilizer (kg ha⁻¹)',
    'corn_yield_Mg': 'Corn grain yield (Mg ha⁻¹)',
    'SOC_0_15cm_Mg': 'Soil organic carbon (0-15cm,Mg ha⁻¹)',
    'Residue_Biomass_Mg': 'Total residue biomass (Mg ha⁻¹)',
    'Incorporated_Biomass_Mg': 'Incorporated Residue biomass (Mg ha⁻¹)',
    'Below_ground_biomass_Mg': 'Below-ground biomass (Mg ha⁻¹)',
    'year': 'Time (Years)',
    'top_carbon_mineralization': 'Soil carbon mineralization (Mg ha⁻¹)',
    'ΔSOC_0_15CM': 'Changes in soil carbon (Mg ha⁻¹) after 105 years',
    'SOC1': 'Changes in soil carbon (Mg ha⁻¹)  after 105 years ',
    '%Δcorn_yield_Mg': "Percentage change in corn grain yield",
    '%ΔResidue_Biomass_Mg': "Percentage change in residue biomass",
    'microbial_carbon': "Microbial biomass carbon (Mg ha⁻¹)"
}


class Plot(PlotManager):
    def __init__(self, data):
        super().__init__()
        self.data = data
        self.plot = plt

    @property
    def results(self):
        return self.data

    def add_report_variable(self):
        pass

    def run(self, **kwargs):
        pass

    def get_simulated_output(self, report_name: str):
        pass


COUNT = {'count': 0}

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading("Seaborn Plots — Report", 0)


def generate_captions(y, x, hue):
    info = {
        'timing': "fertilizer application timing",
        'Residue': "residue incorporation fractions",
        'Nitrogen': LABELS.get('Nitrogen'),
        'ΔSOC_0_15CM': 'Changes in soil organic carbon',
        'corn_yield_Mg': 'Corn grain yield',
        'cnr': "Surface carbon to nitrogen (C:N) ratio",
        "Residue_Biomass_Mg": LABELS.get("Residue_Biomass_Mg"),
        "Surf_Org_Carbon_Mg": LABELS.get('Surf_Org_Carbon_Mg'),
        'Below_ground_biomass_Mg': LABELS.get('Below_ground_biomass_Mg'),
        'SOC_0_15cm_Mg': LABELS.get('SOC_0_15cm_Mg'),
        'Incorporated_Biomass_Mg': 'Incorporated Residue biomass (Mg ha⁻¹)',
        'year': "Time (Years)",
        'top_carbon_mineralization': 'Soil carbon mineralization (Mg ha⁻¹)',
        'top_mineralized_N': 'Mineralized nitrogen (0-15cm, kg ha⁻¹)',
        '%Δcorn_yield_Mg': "Percentage change in corn grain yield"

    }

    return f"{info.get(y)} across different {info.get(x)} and {info.get(hue)}  (colored bars). "


def add_image(filename: str, y, x, hue, save=True, **kwargs):
    COUNT['count'] += 1

    section = doc.sections[0]
    usable_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(filename), width=Inches(usable_width))  # scales to full text width
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph(f"Figure {COUNT['count']}. {generate_captions(y=y, x=x, hue=hue)}")
    try:
        cap.style = "Caption"
    except Exception:
        pass


def relplot(data, show=False, add_scatter=True, legend_loc='upper left', bbox_to_anchor=(0.98, 0.98), **kwargs):
    kwargs.setdefault("aspect", PlotConstants.ASPECT_RATIO)
    kwargs.setdefault("palette", PlotConstants.PALETTE)
    kwargs.setdefault("legend", 'full')
    kwargs.setdefault("height", PlotConstants.HEIGHT)
    kwargs.setdefault("kind", 'line')

    if add_scatter:
        ploter = sns.lineplot
        kwargs.setdefault('markers', True)
        kwargs.pop('kind', None)
        kwargs.pop('aspect', None)
        kwargs.pop('height', None)
    else:
        ploter = sns.relplot
    g = ploter(
        data=data,
        **kwargs
    )
    # plt.setp(g.legend.legend_handles, markeredgewidth=1)

    # rest of your code unchanged
    fileName = RESULTS / f"{kwargs.get('x')}_{kwargs.get('y')}-{kwargs.get('hue')}-{kwargs.get('kind')}-{kwargs.get('height')}{PlotConstants.FILE_EXTENSION}"
    plt.xlabel(LABELS.get(kwargs.get("x", kwargs.get('x'))), fontsize=PlotConstants.xLABEL_SIZE)
    plt.xticks(fontsize=PlotConstants.xTICK_SIZE)
    plt.yticks(fontsize=PlotConstants.yTICK_SIZE)
    y = kwargs.get("y")

    plt.ylabel(LABELS.get(y, y), fontsize=PlotConstants.yLABEL_SIZE)

    plt.tight_layout()
    plt.savefig(fileName)
    if show:
        open_file(fileName)
    add_image(fileName, y=kwargs.get("y"), x=kwargs.get('x'), hue=kwargs.get('hue'), save=True)
    plt.close()
    return g


def cat_plot(data, title="", show=False, legend_loc='upper left', bbox_to_anchor=(0.10, 0.98), **kwargs):
    # turn OFF seaborn legend
    kwargs['legend'] = 'full'
    kwargs.setdefault('width', 0.5)
    # your existing setup...
    if kwargs.get('kind') == 'bar':
        kwargs.pop('showfliers', None)
    elif kwargs.get('kind') == 'box':
        kwargs.setdefault('showfliers', False)
    kwargs.setdefault("height", PlotConstants.HEIGHT)
    kwargs.setdefault("aspect", PlotConstants.ASPECT_RATIO)
    kwargs.setdefault("palette", PlotConstants.PALETTE)

    catPlot = Plot(data=data)
    g = catPlot.cat_plot(**kwargs)  # this returns a seaborn FacetGrid
    sns.move_legend(g, loc=legend_loc, bbox_to_anchor=bbox_to_anchor, title=None)
    # rest of your code unchanged
    fileName = RESULTS / f"{kwargs.get('x')}_{kwargs.get('y')}-{kwargs.get('hue')}-{kwargs.get('kind')}-{kwargs.get('height')}{PlotConstants.FILE_EXTENSION}"
    catPlot.plot.xlabel(LABELS.get(kwargs.get("x")), fontsize=PlotConstants.xLABEL_SIZE)
    catPlot.plot.xticks(fontsize=PlotConstants.xTICK_SIZE)
    catPlot.plot.yticks(fontsize=PlotConstants.yTICK_SIZE)
    y = kwargs.get("y")
    catPlot.plot.ylabel(LABELS.get(y, y), fontsize=PlotConstants.yLABEL_SIZE)
    catPlot.plot.title(title)
    catPlot.plot.tight_layout()
    catPlot.plot.savefig(fileName)
    if show:
        open_file(fileName)
    add_image(fileName, y=kwargs.get("y"), x=kwargs.get('x'), hue=kwargs.get('hue'), save=True)
    catPlot.plot.close()
    plt.close()


def mva_plot(data, **kwargs):
    kwargs.setdefault("height", PlotConstants.HEIGHT)
    kwargs.setdefault('time_col', 'year')
    kwargs.setdefault('table', data)
    kwargs.setdefault('response', kwargs.get('y'))
    g = plot.plot_mva(table=da, response='SOC_0_15cm_Mg', time_col='year', hue='Residue', palette='tab10',
                      height=8, aspect=1.5, errorbar=None, col='Nitrogen', col_wrap=2, )
    base = Plot(data=data)
    kwargs.setdefault('ylabel', LABELS.get('response'))
    kwargs.setdefault('palette', PlotConstants.PALETTE)
    base.plot_mva(**kwargs)
    y = kwargs.get('response')
    g.set_axis_labels("", "")
    xtick_size = PlotConstants.xTICK_SIZE
    ytick_size = PlotConstants.yTICK_SIZE
    ylabel_size = PlotConstants.yLABEL_SIZE
    xlabel_size = PlotConstants.xLABEL_SIZE
    ylabel = LABELS.get(y)
    xlabel = LABELS.get(kwargs.get('time_col', kwargs.get('x')))
    # Enforce tick sizes for every facet (robustly)
    for ax in g.axes.flat:
        ax.tick_params(axis='x', which='both', labelsize=xtick_size)
        ax.tick_params(axis='y', which='both', labelsize=ytick_size)
        # Fallback in case styles/backends override tick_params
        plt.setp(ax.get_xticklabels(), fontsize=xtick_size)
        plt.setp(ax.get_yticklabels(), fontsize=ytick_size)
        # Also scale the scientific-notation offset text if present
        ax.yaxis.get_offset_text().set_size(ytick_size)

    # Shared labels
    g.fig.supylabel(ylabel, x=0.002, fontsize=ylabel_size)
    g.fig.supxlabel(xlabel, y=0.002, fontsize=xlabel_size)

    # Legend cleanup + size
    leg = getattr(g, "_legend", None) or getattr(g, "legend", None)
    if leg is not None:
        leg.set_title(None)
        for txt in leg.get_texts():
            txt.set_fontsize(min(xtick_size, ytick_size))

    base.plot.tight_layout()

    fileName = RESULTS / f"_one_file{kwargs.get("time_col")}_{kwargs.get("y")}-{kwargs.get('hue')}-{kwargs.get('response')}-{kwargs.get('height')}{PlotConstants.FILE_EXTENSION}"
    base.plot.savefig(fileName, dpi=PlotConstants.DPI)
    open_file(fileName)
    plt.close()
    base.plot.close()


if __name__ == '__main__':
    data = merge_tables(data=RESULTS / 'one_single_split_single.db', tables=['carbon', 'yield'])
    sin = merge_tables(str(RESULTS / 'single_model.db'), tables=['carbon', 'yield'])
    da = data.copy()
    # sin = sin.assign(timing='Single')
    # sin.eval('ce = top_carbon_mineralization/SOC_0_15cm_Mg *top_carbon_mineralization', inplace=True)
    split = merge_tables(str(RESULTS / 'split_model.db'), tables=['carbon', 'yield'])
    # split = split.assign(timing='Split')
    # split.eval('ce = SOC_0_15cm_Mg/Incorporated_Biomass_Mg *SOC_0_15cm_Mg', inplace=True)
    # da = pd.concat([sin, split], ignore_index=True)


    plot = Plot(data=da)
    da.eval('CFE =(SOC_0_15cm_Mg/Incorporated_Biomass_Mg)* Incorporated_Biomass_Mg', inplace=True)

    mva_plot(data=data, response='SOC_0_15cm_Mg', grouping=("Nitrogen", 'Residue',), hue='Residue',
             errorbar=None,
             col="Residue", col_wrap=2)
    mva_plot(data=da, response='grainwt', grouping=("Nitrogen", 'Residue', ), hue='Residue',
             errorbar=None,
             col="Nitrogen", col_wrap=2)
    # keep only LABELS that exist in da
    cols = [c for c in LABELS.keys() if c in da.columns]

    # keep only numeric among those
    num_cols = da[cols].select_dtypes(include="number").columns.tolist()

    # compute means, keeping NaN groups if you want
    # dat = (
    #     da.groupby(['Nitrogen', 'Residue', 'timing'], observed=True)[num_cols]
    #     .mean()  # pandas ≥2.0 auto numeric_only
    #     .reset_index()
    # )
    dat = da

    cat_plot(data=dat, y='Incorporated_Biomass_Mg', x='Nitrogen', hue='Nitrogen', kind='bar', title='')

    cat_plot(data=dat, y='Surf_Org_Carbon_Mg', x='Nitrogen', kind='box', title='', n_boot=500,
             legend='full', legend_out=True)
    dt = dat.groupby(['Nitrogen', 'Residue', ], observed=True)['Surf_Org_Carbon_Mg'].mean().reset_index()
    cat_plot(data=dt, y='Surf_Org_Carbon_Mg', x='Residue',  kind='bar', title='',
             legend='full', legend_out=True, legend_loc='upper right', bbox_to_anchor=(0.98, 0.98))

    cat_plot(data=dat, y='SOC_0_15cm_Mg', x='Nitrogen', hue='Residue', kind='bar', title='')
    cat_plot(data=dat, y='cnr', x='Nitrogen',  kind='box', title='', estimator='mean',
             legend_loc='upper right', bbox_to_anchor=(0.98, 0.98), )
    cat_plot(data=dat, y='cnr', x='Residue', hue='timing', kind='bar', title='', estimator='mean',
             legend_loc='upper right', bbox_to_anchor=(0.98, 0.98), )

    cat_plot(data=dat, y='top_carbon_mineralization', x='Nitrogen', hue='timing', kind='box', title='',
             estimator='mean', legend_loc='upper right', bbox_to_anchor=(0.98, 0.98), )
    cat_plot(data=dat, y='Residue_Biomass_Mg', x='Nitrogen', hue='timing', kind='box', title='',
             estimator='mean', legend_loc='upper right', bbox_to_anchor=(0.98, 0.98), )
    cat_plot(data=dat, y='microbial_carbon', x='Residue', kind='bar', title='',
             estimator='mean', legend_loc='upper right', bbox_to_anchor=(0.98, 0.98), )
    single = da[da['timing'] == '1']
    cat_plot(data=da, show=False, y='microbial_carbon', x='Nitrogen', hue='Residue', kind='box', title='', n_boot=1000,
             col='timing', col_wrap=1, height=6, aspect=2,
             estimator='mean', legend_loc='upper left', bbox_to_anchor=(0.1, 1), )

    soc_change = calculate_soc_changes(data=dat, col='SOC_0_15CM', grouping=('Nitrogen', "year", "Residue"))
    print('soc changes:\n==========', soc_change)

    from change_metrics import compute_last_minus_first_change, mean_diff_between_windows

    c_change = compute_last_minus_first_change(data=da, col='SOC_0_15CM', grouping=('Nitrogen', 'Residue', 'timing'))
    cat_plot(data=c_change, y='ΔSOC_0_15CM', x='Nitrogen', kind='box', title='',
             estimator='mean', legend_loc='upper right', bbox_to_anchor=(0.98, 0.98), )

    w = 10
    year_start = 1904
    year_ending = 2005
    first_window = year_start, year_start + w
    last_window = year_ending - w, year_ending

    # ________________grain and total biomass changes ___________________________________________________________

    dif_grain = mean_diff_between_windows(da, grouping=['Nitrogen', 'Residue', ], sample=False, sample_size=10,
                                          sample_window=20,
                                          col='corn_yield_Mg', last_window=last_window, first_window=first_window,
                                          min_obs=5)
    cat_plot(data=dif_grain, show=True, y='%Δcorn_yield_Mg', x='Nitrogen', kind='box', title='',
             estimator='mean', legend_loc='upper right', bbox_to_anchor=(0.2, 0.98), )
    relplot(data=dif_grain, show=True, x='Nitrogen', y='%Δcorn_yield_Mg', hue='Residue', kind='line', add_scatter=False)

    change_reb = mean_diff_between_windows(da, grouping=['Nitrogen', 'Residue', ], sample=True, sample_window=20,
                                           sample_size=5,
                                           col='Residue_Biomass_Mg', last_window=last_window, first_window=first_window,
                                           min_obs=5)
    mic_soc15 = mean_diff_between_windows(da, grouping=['Nitrogen', 'Residue'], sample=True, sample_window=20,
                                          col='microbial_carbon', last_window=last_window, first_window=first_window,
                                          min_obs=5)

    cat_plot(data=mic_soc15, show=False, y='%Δmicrobial_carbon', x='Residue', hue='Nitrogen', kind='bar', title='',
             estimator='mean', legend_loc='upper right', bbox_to_anchor=(0.16, 0.98), col='timing', col_wrap=1,
             aspect=2)

    c_change = compute_last_minus_first_change(data=da, col='SOC_0_15CM', grouping=('Nitrogen', 'Residue'))

    c_change = compute_last_minus_first_change(data=da, col='SOC_0_15CM', grouping=('Nitrogen', 'Residue'))

    c_change = compute_last_minus_first_change(data=da, col='SOC_0_15CM', grouping=('Nitrogen', 'Residue'))
    cat_plot(data=c_change, y='ΔSOC_0_15CM', x='Nitrogen', hue='timing', kind='bar', title='',
             estimator='mean', legend_loc='upper right', bbox_to_anchor=(0.2, 0.98), )

    doc_path = RESULTS / f"seaborn_plots_report_from{Path(__file__).name}.docx"

    from stat_module import dynamic_paired_ttest, dynamic_independent_ttest


    def select_from_data(data, Residue='1', Nitrogen='326', col='grainwt'):
        da = data.copy()
        df = da.loc[(da['Nitrogen'].astype(str) == str(Nitrogen)) & (da['Residue'].astype(str) == str(Residue))]

        s = df[df['timing'].astype(str) == '1'][col]

        sp = df[df['timing'].astype(str) != '1'][col]
        assert not sp.empty
        assert not s.empty
        return s, sp


    dynamic_independent_ttest(*select_from_data(da, Residue='0.75', Nitrogen='202', col='SurfaceOrganicMatter_Carbon'))

    dynamic_paired_ttest(res[f"corn_yield_Mg@{'-'.join(map(str, last_window))}"],
                         res[f"corn_yield_Mg@{'-'.join(map(str, first_window))}"])

    from utils import find_soc_equilibrium

    seq = find_soc_equilibrium(da, group_cols=("Nitrogen", "Residue", 'timing'), soc_col="SOC_0_15CM", year_col="year",
                               min_stable_years=3, window=7)

    g = sns.relplot(data=seq, x='Nitrogen', y='equilibrium_year', kind='line', style='timing', col='Residue',
                    col_wrap=2)

    fin = RESULTS / '_peaks_.png'
    plt.savefig(fin, dpi=600)
    open_file(fin)
    plt.close()

    from scipy.stats import pearsonr

    da.eval("RXN =R*N", inplace=True)
    gcol = 'Nitrogen'
    result = (
        da.groupby([gcol], observed=False)
        .apply(lambda g: pd.Series(pearsonr(g['RXN'], g['grainwt']),
                                   index=['r', 'pvalue'])).reset_index()
    )

    print(result)
    cat_plot(data=result, title='', show=True, x=gcol, y='r', kind='bar', legend_loc='upper right')
    from scipy.stats import linregress

    import pandas as pd
    import matplotlib.pyplot as plt
    from statsmodels.graphics.factorplots import interaction_plot
    import numpy as np
    import pandas as pd
    from scipy.stats import linregress


    # Ensure categorical labels (especially if Nitrogen is numeric)

    def _linreg_group(g: pd.DataFrame) -> pd.Series:
        # ensure numeric, drop NaNs pairwise
        x = pd.to_numeric(g['year'], errors='coerce')
        y = pd.to_numeric(g['grainwt'], errors='coerce')
        m = x.notna() & y.notna()
        if m.sum() < 2:
            # not enough data to regress; return NaNs
            return pd.Series({'slope': np.nan, 'intercept': np.nan,
                              'rvalue': np.nan, 'pvalue': np.nan,
                              'stderr': np.nan, 'intercept_stderr': np.nan})
        res = linregress(x[m].to_numpy(), y[m].to_numpy())
        return pd.Series(res._asdict())  # names: slope, intercept, rvalue, pvalue, stderr, intercept_stderr


    st = da

    relplot(data=st, show=True, x='Nitrogen', y='cnr', hue='Residue', linewidth=1.5,
            kind='line', errorbar=None, style='Residue', markers='o', )

    relplot(data=st, show=True, x='Nitrogen', y='corn_yield_Mg', hue='Residue', linewidth=2,
            kind='line', errorbar=None, style='Residue', markers='o', height=12, aspect=1.5)
    relplot(data=st, show=True, x='Nitrogen', y='Below_ground_biomass_Mg', hue='Residue', estimator='mean',
            kind='line', errorbar=None, markers='o', linewidth=1.3)
    doc.save(doc_path)

    import statsmodels.formula.api as smf

    # Fit a quadratic (second-order) interaction model

    data = da[da['timing'] != '1']
    res_col = 'corn_yield_Mg'
    data = data.groupby(['R', "N"])[res_col].mean().reset_index()
    model_quad = smf.ols(
        f'{res_col} ~ N + R + I(N**2) + I(R**2) + N:R',
        data=data
    ).fit()

    print(model_quad.summary())

    b = model_quad.params
    N_opt = (2 * b['I(R ** 2)'] * b['N'] - b['R'] * b['N:R']) / \
            (b['N:R'] ** 2 - 4 * b['I(N ** 2)'] * b['I(R ** 2)'])
    R_opt = (-b['R'] - 2 * b['I(R ** 2)'] * N_opt - b['N:R'] * N_opt) / (2 * b['I(R ** 2)'])
    print(f"Optimal Nitrogen ≈ {N_opt:.2f}, Residue ≈ {R_opt:.2f}")

    # Observed and predicted
    y_true = model_quad.model.endog
    y_pred = model_quad.fittedvalues

    # RMSE
    rmse = np.sqrt(np.mean((y_true - y_pred) ** 2))
    print(f"RMSE: {rmse:.3f}")

    import numpy as np
    from sklearn.metrics import mean_squared_error, r2_score

    # --- RMSE
    rmse = np.sqrt(mean_squared_error(y_true, y_pred))

    # --- RRMSE (%)
    rrmse = (rmse / np.mean(y_true)) * 100

    # --- R² (Coefficient of Determination)
    r2 = r2_score(y_true, y_pred)

    print(f"RMSE  : {rmse:.3f}")
    print(f"RRMSE : {rrmse:.2f}%")
    print(f"R²     : {r2:.4f}")
